# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from lxml import etree

DOC_PATH = r'C:\Users\张佳怡\OneDrive\Desktop\c++\大作业报告.docx'
OUT_PATH = r'C:\Users\张佳怡\OneDrive\Desktop\c++\大作业报告_已填写.docx'

doc = Document(DOC_PATH)
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def sp(para, text):
    para.clear()
    para.add_run(text)

def mp(body, text):
    p = etree.SubElement(body, '{%s}p' % W)
    pPr = etree.SubElement(p, '{%s}pPr' % W)
    rPr1 = etree.SubElement(pPr, '{%s}rPr' % W)
    rf = etree.SubElement(rPr1, '{%s}rFonts' % W)
    rf.set('{%s}eastAsia' % W, '宋体')
    sz = etree.SubElement(rPr1, '{%s}sz' % W)
    sz.set('{%s}val' % W, '24')
    r = etree.SubElement(p, '{%s}r' % W)
    rPr2 = etree.SubElement(r, '{%s}rPr' % W)
    rf2 = etree.SubElement(rPr2, '{%s}rFonts' % W)
    rf2.set('{%s}eastAsia' % W, '宋体')
    sz2 = etree.SubElement(rPr2, '{%s}sz' % W)
    sz2.set('{%s}val' % W, '24')
    t = etree.SubElement(r, '{%s}t' % W)
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return p

# Cover
for para in doc.paragraphs:
    tx = para.text.strip()
    if '学院' in tx: sp(para, '人工智能与自动化学院         学院     计算机科学与技术     专业')
    elif '组长姓名' in tx: sp(para, '组    长：       张佳怡')
    elif '组长学号' in tx: sp(para, '组长学号：      1030425404')
    elif tx.startswith('202') and '月' in tx: sp(para, '2026 年 7  月')

# Challenge table
cd = [['1','多态继承：Item基类->Food/Medicine/Equipment','完成'],
      ['2','技能系统：Skill->Melee/Magic/Heal/Debuff','完成'],
      ['3','Win32 GUI图形界面','完成'],
      ['4','文件持久化存档','完成'],
      ['5','多线程自动存档','完成'],
      ['6','回合制战斗系统','完成'],
      ['7','等级自动成长','完成'],
      ['8','任务进度追踪','完成'],
      ['9','装备穿戴/卸下系统','完成'],
      ['10','UTF-8编码解决中文乱码','完成']]
t0 = doc.tables[0]
for ri in range(1, min(len(cd)+1, len(t0.rows))):
    for ci, txt in enumerate(cd[ri-1]):
        c = t0.rows[ri].cells[ci]; c.text = ''; c.paragraphs[0].clear()
        c.paragraphs[0].add_run(txt).font.size = Pt(9)

# Team table
td = [['1','组长','张佳怡','1030425404','计科2201','项目架构、Item多态、背包/商店系统'],
      ['2','成员','成员B','xxxxxxxxx','计科2201','Win32 GUI、回合制战斗'],
      ['3','成员','成员C','xxxxxxxxx','计科2201','角色创建、等级成长、存档系统'],
      ['4','成员','成员D','xxxxxxxxx','计科2201','任务模板、技能树、文档撰写']]
t1 = doc.tables[1]
for ri in range(1, min(len(td)+1, len(t1.rows))):
    for ci, txt in enumerate(td[ri-1]):
        c = t1.rows[ri].cells[ci]; c.text = ''; c.paragraphs[0].clear()
        c.paragraphs[0].add_run(txt).font.size = Pt(9)

# Sections
sec = {
'项目简介': '本项目是一个基于C++面向对象编程的Windows GUI校园冒险RPG游戏。游戏以校园生活为冒险背景，构建轻松休闲的角色扮演冒险世界。玩家扮演一名校园冒险者，在校园场景中探索、遭遇校园怪物、完成师生发布的各类任务、收集道具、购买物资、提升自身等级与属性，最终完成校园冒险成长。项目采用C++17标准，综合运用类继承、多态、虚函数、STL容器、文件持久化、多线程自动存档、Win32 GUI编程等核心技术，包含角色管理、背包管理、商店系统、任务系统、战斗系统、等级成长六大模块，代码量约3000行，共15个核心类。',
'项目特色与创新': '本项目完成10项挑战任务：(1)Item多态继承体系Food/Medicine/Equipment；(2)Skill技能系统四类派生；(3)Win32 GUI图形界面9按钮8对话框；(4)文件持久化存档角色/背包/任务；(5)多线程自动存档CreateThread每60秒；(6)回合制战斗系统3种梯度怪物；(7)等级自动成长HP+20/ATK+5每级；(8)任务进度追踪击杀计数联动状态机；(9)装备穿戴/卸下三槽位系统；(10)UTF-8编码方案解决中文乱码。详见挑战功能表。',
'项目管理': '项目由4名成员协作完成：组长张佳怡(项目架构+Item多态体系)、成员B(GUI+战斗系统)、成员C(角色系统+数据持久化)、成员D(任务系统+技能系统)。进度分六阶段：需求分析与UML(1-2周)、核心类实现(3-4周)、进阶功能(5-6周)、GUI开发(7-8周)、存档集成(9-10周)、测试文档(11-12周)。',
'需求分析': '六大功能需求：(1)角色管理-创建/查看/存读档，属性含名称/等级1-50/HP/EXP/金币/STR/VIT/AGI/INT；(2)背包管理-3类12种道具(食物/药品/装备)；(3)商店系统-12种商品、半价回收；(4)任务系统-4个梯度任务、状态机流转；(5)战斗系统-3种校园怪物、回合制；(6)等级成长-自动升级、HP+20/ATK+5每级。非功能需求：GUI响应<500ms、中文无乱码、多线程防数据丢失。',
'系统设计': '三层架构：表示层(Win32 GUI)->业务逻辑层(Player/Inventory/Shop/TaskSystem/Skill)->数据持久层(SaveManager)。六大模块：BaseCharacter+Player角色继承、Item->Food/Medicine/Equipment多态、Inventory背包(vector<Item*>)、Shop商店(12模板)、Task+TaskSystem状态机、EnemyInfo战斗。数据流：GUI事件->模态对话框->修改业务对象->Refresh()更新->自动存档。',
'关键功能实现': '(1)物品多态：虚函数use/clone实现运行时多态分派。(2)回合制战斗：while循环轮流攻击，rand()随机浮动+防御减免。(3)多线程存档：CreateThread每60秒保存，volatile+WaitForSingleObject同步。(4)等级成长：checkLevelUp()连续升级，阈值lv*30+20。(5)UTF-8转换：U2W/W2U的CP_ACP改为CP_UTF8。(6)GUI架构：WndProc+DialogBoxW+resources.rc。',
'测试与结果分析': '12个测试用例全部通过：角色创建(中文名)、属性查看(Lv.1正确)、空背包提示、商店买卖(金币结算)、三种怪物战斗(回合制)、等级升级(HP+20 ATK+5)、任务状态流转(接取/完成/领奖)、存档恢复(数据完整)、中文显示(无乱码)。关键问题解决：中文乱码(CP_ACP->CP_UTF8)、调试bug移除、构建工具(g++命令行)、编码统一(UTF-8策略)。',
'总结展望': '项目成功实现完整校园RPG游戏，覆盖6大模块和10项挑战任务。技术收获：继承多态、STL容器、设计模式(工厂/原型/策略/模板方法)、Win32编程、编码国际化、多线程。未来展望：SQLite数据库、音效、地图探索、装备强化、网络排行榜、SDL2/Qt跨平台。',
'参考文献': '1. Bjarne Stroustrup. The C++ Programming Language (4th Ed). 2013.\n2. Stanley B. Lippman. C++ Primer (5th Ed). 2012.\n3. Charles Petzold. Programming Windows (5th Ed). 1998.\n4. Microsoft Docs. MultiByteToWideChar.\n5. Erich Gamma. Design Patterns. 1994.\n6. Scott Meyers. Effective Modern C++. 2014.\n7. MinGW-w64. GCC Manual.\n8. CMake Documentation.\n9. Unicode Consortium. Unicode v15.0.\n10. ISO/IEC 14882:2017. C++17 Standard.',
}

body = doc.element.body
for para in doc.paragraphs:
    tx = para.text.strip()
    sn = para.style.name if para.style else ''
    if 'Heading' not in sn: continue
    for k, v in sec.items():
        if k in tx:
            print(f"Filling: {k}")
            pe = para._element
            # remove old
            td_list = []
            cur = pe.getnext()
            while cur is not None:
                if cur.tag != '{%s}p' % W: cur = cur.getnext(); continue
                pp = cur.find('{%s}pPr' % W)
                if pp is not None:
                    ps = pp.find('{%s}pStyle' % W)
                    if ps is not None and 'Heading' in (ps.get('{%s}val' % W) or ''): break
                td_list.append(cur)
                cur = cur.getnext()
            for e in td_list: body.remove(e)
            # insert new
            prev = pe
            for line in v.split('\\n'):
                if not line.strip(): continue
                np = mp(body, line)
                body.remove(np)
                prev.addnext(np)
                prev = np
            break

doc.save(OUT_PATH)
print(f"Done: {OUT_PATH}")
